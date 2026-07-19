import os
import re
import sqlite3
import sys
import webbrowser
import ctypes
import ctypes.wintypes
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PySide6.QtCore import QDate, QSettings, Qt
from PySide6.QtGui import QAction, QColor, QIcon, QPalette, QBrush
from PySide6.QtWidgets import (
    QApplication,
    QCheckBox,
    QDateEdit,
    QDialog,
    QDialogButtonBox,
    QColorDialog,
    QFrame,
    QFormLayout,
    QGridLayout,
    QHBoxLayout,
    QHeaderView,
    QInputDialog,
    QLabel,
    QLineEdit,
    QMainWindow,
    QMenu,
    QMessageBox,
    QPushButton,
    QPlainTextEdit,
    QScrollArea,
    QSizeGrip,
    QTableWidget,
    QTableWidgetItem,
    QTabWidget,
    QToolBar,
    QVBoxLayout,
    QWidget,
)


APP_NAME = "DailyNotebook"
APP_DISPLAY_NAME = "林的记事本"
PROJECT_DIR = Path(__file__).resolve().parent
DB_PATH = PROJECT_DIR / "tasks.db"
JOURNAL_DOCX_PATH = PROJECT_DIR / "life_journal.docx"
ICON_PATH = PROJECT_DIR / "daily_notebook.ico"
RUN_KEY = r"Software\Microsoft\Windows\CurrentVersion\Run"
WORD_FONT = "SimSun"
WORD_FONT_EAST_ASIA = "SimSun"
CHATGPT_URL = "https://chatgpt.com/"

TRANSPARENT_STYLE = """
QMainWindow {
    background: transparent;
}
#RootPanel {
    background: rgba(246, 248, 252, 84);
    border: 1px solid rgba(255, 255, 255, 155);
    border-radius: 12px;
}
#TitleBar {
    background: rgba(245, 247, 250, 132);
    border-top-left-radius: 12px;
    border-top-right-radius: 12px;
}
#TitleButton {
    color: rgb(14, 18, 28);
    background: transparent;
    border: 0;
    border-radius: 5px;
    padding: 2px 9px;
}
#TitleButton:hover {
    background: rgba(255, 255, 255, 160);
}
QToolBar {
    background: rgba(245, 247, 250, 132);
    border: 0;
    border-bottom: 1px solid rgba(20, 24, 34, 42);
    spacing: 4px;
    padding: 4px;
}
QToolButton {
    color: rgb(14, 18, 28);
    background: rgba(255, 255, 255, 158);
    border: 1px solid rgba(20, 24, 34, 64);
    border-radius: 5px;
    padding: 4px 7px;
}
QToolButton:hover {
    background: rgba(255, 255, 255, 210);
}
QToolButton:pressed {
    background: rgba(218, 226, 240, 220);
}
QPushButton {
    color: rgb(14, 18, 28);
    background: rgba(255, 255, 255, 158);
    border: 1px solid rgba(20, 24, 34, 64);
    border-radius: 5px;
    padding: 4px 9px;
}
QPushButton:hover {
    background: rgba(255, 255, 255, 210);
}
QPushButton:pressed {
    background: rgba(218, 226, 240, 220);
}
QMenu {
    color: rgb(14, 18, 28);
    background: rgba(255, 255, 255, 242);
    border: 1px solid rgba(20, 24, 34, 80);
}
QMenu::item {
    padding: 6px 24px;
}
QMenu::item:selected {
    color: white;
    background: rgba(47, 93, 158, 210);
}
QTabWidget::pane {
    border: 1px solid rgba(20, 24, 34, 52);
    background: rgba(255, 255, 255, 58);
    border-radius: 6px;
}
QTabWidget, QStackedWidget, QStackedWidget > QWidget {
    background: transparent;
}
QTabBar::tab {
    color: rgb(14, 18, 28);
    background: rgba(255, 255, 255, 124);
    border: 1px solid rgba(20, 24, 34, 44);
    padding: 5px 12px;
    margin-right: 2px;
}
QTabBar::tab:selected {
    color: white;
    background: rgba(47, 93, 158, 195);
}
QTableWidget {
    color: rgb(10, 14, 24);
    background: transparent;
    alternate-background-color: transparent;
    gridline-color: rgba(20, 24, 34, 48);
    selection-background-color: rgba(47, 93, 158, 165);
    selection-color: white;
    border: 0;
}
QAbstractScrollArea {
    background: transparent;
}
QAbstractScrollArea > QWidget {
    background: transparent;
}
QTableWidget::viewport {
    background: transparent;
}
QTableWidget::item {
    color: rgb(10, 14, 24);
    padding: 4px;
}
QHeaderView::section {
    color: rgb(10, 14, 24);
    background: rgba(255, 255, 255, 162);
    border: 0;
    border-right: 1px solid rgba(20, 24, 34, 42);
    border-bottom: 1px solid rgba(20, 24, 34, 52);
    padding: 4px;
}
QTableCornerButton::section {
    background: rgba(255, 255, 255, 162);
    border: 0;
}
QCheckBox {
    color: rgb(10, 14, 24);
    spacing: 5px;
}
QCheckBox::indicator {
    width: 15px;
    height: 15px;
}
QLabel {
    color: rgb(10, 14, 24);
}
QScrollBar:vertical, QScrollBar:horizontal {
    background: rgba(255, 255, 255, 70);
    border: 0;
    width: 10px;
    height: 10px;
}
QScrollBar::handle:vertical, QScrollBar::handle:horizontal {
    background: rgba(20, 24, 34, 95);
    border-radius: 5px;
}
QScrollBar::add-line, QScrollBar::sub-line {
    width: 0;
    height: 0;
}
"""


def build_transparent_style(text_color, text_opacity, background_opacity):
    alpha = max(20, min(230, int(255 * background_opacity / 100)))
    text_alpha = max(30, min(255, int(255 * text_opacity / 100)))
    text_rgb = f"rgba({text_color.red()}, {text_color.green()}, {text_color.blue()}, {text_alpha})"
    bg_rgb = "246, 248, 252"
    style = TRANSPARENT_STYLE
    style = style.replace("rgb(10, 14, 24)", text_rgb)
    style = style.replace("rgb(14, 18, 28)", text_rgb)
    replacements = {
        "rgba(246, 248, 252, 84)": f"rgba({bg_rgb}, {alpha})",
        "rgba(245, 247, 250, 132)": f"rgba({bg_rgb}, {min(240, alpha + 45)})",
        "rgba(255, 255, 255, 158)": f"rgba({bg_rgb}, {min(245, alpha + 70)})",
        "rgba(255, 255, 255, 124)": f"rgba({bg_rgb}, {min(235, alpha + 45)})",
        "rgba(255, 255, 255, 58)": f"rgba({bg_rgb}, {max(18, int(alpha * 0.72))})",
        "rgba(255, 255, 255, 42)": f"rgba({bg_rgb}, {max(16, int(alpha * 0.58))})",
        "rgba(255, 255, 255, 28)": f"rgba({bg_rgb}, {max(10, int(alpha * 0.42))})",
        "rgba(255, 255, 255, 70)": f"rgba({bg_rgb}, {max(22, int(alpha * 0.82))})",
        "rgba(255, 255, 255, 162)": f"rgba({bg_rgb}, {min(248, alpha + 85)})",
    }
    for old, new in replacements.items():
        style = style.replace(old, new)
    return style


def now_text():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def parse_date(value):
    return datetime.strptime(value, "%Y-%m-%d").date()


def duration_text(start_time, end_time):
    if not start_time or not end_time:
        return ""
    start = datetime.strptime(start_time, "%Y-%m-%d %H:%M:%S")
    end = datetime.strptime(end_time, "%Y-%m-%d %H:%M:%S")
    seconds = max(0, int((end - start).total_seconds()))
    hours, rem = divmod(seconds, 3600)
    minutes, _ = divmod(rem, 60)
    if hours:
        return f"{hours}小时{minutes}分钟"
    return f"{minutes}分钟"


def apply_word_font(run, size=9, bold=False):
    run.font.name = WORD_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), WORD_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), WORD_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT_EAST_ASIA)
    run._element.rPr.rFonts.set(qn("w:cs"), WORD_FONT)


EMOJI_RE = re.compile(
    "["
    "\U0001F1E6-\U0001F1FF"
    "\U0001F300-\U0001F5FF"
    "\U0001F600-\U0001F64F"
    "\U0001F680-\U0001F6FF"
    "\U0001F700-\U0001F77F"
    "\U0001F780-\U0001F7FF"
    "\U0001F800-\U0001F8FF"
    "\U0001F900-\U0001F9FF"
    "\U0001FA00-\U0001FAFF"
    "\uD800-\uDFFF"
    "\u2600-\u27BF"
    "]+"
)


def clean_chatgpt_response(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = EMOJI_RE.sub("", text)
    cleaned_lines = []
    in_code_block = False
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
        line = re.sub(r"`([^`]*)`", r"\1", line)
        line = re.sub(r"^\s*[-*+]\s+", "- ", line)
        line = re.sub(r"^\s*[>\|]\s*", "> ", line)
        line = re.sub(r"^\d+[\.\)\u3001\uff09]\s*", "", line)
        line = re.sub(r"^#+\s*", "", line).strip()
        line = line.replace("*", "").strip()
        line = re.sub(r"^[\u200b\s]+", "", line).strip()
        if line:
            cleaned_lines.append(line)
    return "\n".join(_split_response_line(line) for line in cleaned_lines if line).strip()


def _split_response_line(line):
    prefix = ""
    content = line
    if line.startswith("- "):
        prefix = "- "
        content = line[2:].strip()
    elif line.startswith("> "):
        prefix = "> "
        content = line[2:].strip()

    content = re.sub(r"([\u3002\uff01\uff1f!?][\"'\u201d\u2019]?)\s*", r"\1\n", content)
    content = re.sub(r"(?<!\d)\.(?=\s+|$)", ".\n", content)
    content = re.sub(r"\n+", "\n", content).strip()
    parts = content.splitlines()
    parts = [part.strip() for part in parts if part.strip()]
    if not parts:
        return prefix.rstrip()
    return "\n".join(prefix + part for part in parts)


class TaskStore:
    def __init__(self, path):
        self.path = path
        self.conn = sqlite3.connect(path)
        self.conn.row_factory = sqlite3.Row
        self.init_schema()

    def init_schema(self):
        self.conn.execute(
            """
            CREATE TABLE IF NOT EXISTS tasks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                notes TEXT DEFAULT '',
                due_date TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'pending',
                start_time TEXT,
                end_time TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                logged_at TEXT
            )
            """
        )
        self.conn.execute(
            """
            CREATE TABLE IF NOT EXISTS essays (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT DEFAULT '',
                content TEXT NOT NULL,
                ai_response TEXT DEFAULT '',
                start_time TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            """
        )
        columns = {
            row["name"]
            for row in self.conn.execute("PRAGMA table_info(essays)").fetchall()
        }
        if "ai_response" not in columns:
            self.conn.execute("ALTER TABLE essays ADD COLUMN ai_response TEXT DEFAULT ''")
        self.conn.commit()

    def add_task(self, title, notes, due_date):
        stamp = now_text()
        self.conn.execute(
            """
            INSERT INTO tasks (title, notes, due_date, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (title, notes, due_date, stamp, stamp),
        )
        self.conn.commit()

    def update_task(self, task_id, title, notes, due_date):
        self.conn.execute(
            """
            UPDATE tasks
            SET title = ?, notes = ?, due_date = ?, updated_at = ?
            WHERE id = ?
            """,
            (title, notes, due_date, now_text(), task_id),
        )
        self.conn.commit()

    def delete_task(self, task_id):
        self.conn.execute("DELETE FROM tasks WHERE id = ?", (task_id,))
        self.conn.commit()

    def start_task(self, task_id):
        self.conn.execute(
            """
            UPDATE tasks
            SET status = 'running', start_time = COALESCE(start_time, ?), updated_at = ?
            WHERE id = ? AND status != 'done'
            """,
            (now_text(), now_text(), task_id),
        )
        self.conn.commit()

    def complete_task(self, task_id):
        row = self.get_task(task_id)
        if row is None:
            return None
        if row["status"] == "done":
            return row
        stamp = now_text()
        start_time = row["start_time"] or stamp
        self.conn.execute(
            """
            UPDATE tasks
            SET status = 'done', start_time = ?, end_time = ?, updated_at = ?
            WHERE id = ?
            """,
            (start_time, stamp, stamp, task_id),
        )
        self.conn.commit()
        return self.get_task(task_id)

    def mark_logged(self, task_id):
        self.conn.execute(
            "UPDATE tasks SET logged_at = ?, updated_at = ? WHERE id = ?",
            (now_text(), now_text(), task_id),
        )
        self.conn.commit()

    def get_task(self, task_id):
        return self.conn.execute("SELECT * FROM tasks WHERE id = ?", (task_id,)).fetchone()

    def list_tasks(self):
        return self.conn.execute(
            """
            SELECT *
            FROM tasks
            ORDER BY
                CASE status WHEN 'running' THEN 0 WHEN 'pending' THEN 1 ELSE 2 END,
                due_date ASC,
                created_at ASC
            """
        ).fetchall()

    def add_essay(self, title, content, ai_response, start_time):
        stamp = now_text()
        self.conn.execute(
            """
            INSERT INTO essays (title, content, ai_response, start_time, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (title, content, ai_response, start_time, stamp, stamp),
        )
        self.conn.commit()

    def update_essay(self, essay_id, title, content, ai_response):
        self.conn.execute(
            """
            UPDATE essays
            SET title = ?, content = ?, ai_response = ?, updated_at = ?
            WHERE id = ?
            """,
            (title, content, ai_response, now_text(), essay_id),
        )
        self.conn.commit()

    def delete_essay(self, essay_id):
        self.conn.execute("DELETE FROM essays WHERE id = ?", (essay_id,))
        self.conn.commit()

    def get_essay(self, essay_id):
        return self.conn.execute("SELECT * FROM essays WHERE id = ?", (essay_id,)).fetchone()

    def list_essays(self):
        return self.conn.execute(
            """
            SELECT *
            FROM essays
            ORDER BY start_time DESC, id DESC
            """
        ).fetchall()

    def close(self):
        self.conn.close()


class JournalWordLogger:
    def __init__(self, path):
        self.path = path

    def rewrite(self, tasks, essays):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("生活记录")
        apply_word_font(run, size=16, bold=True)

        subtitle = doc.add_paragraph("按时间记录完成的事情、当时的想法、感受和感悟。")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(16)
        for run in subtitle.runs:
            apply_word_font(run, size=9)

        entries = []
        for task in tasks:
            if task["status"] == "done" and task["end_time"]:
                entries.append((task["end_time"], "task", task))
        for essay in essays:
            entries.append((essay["start_time"], "essay", essay))
        entries.sort(key=lambda item: (item[0], 0 if item[1] == "task" else 1))

        for index, (_, entry_type, entry) in enumerate(entries):
            if index:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)

            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
            heading_text = self._heading_text(entry_type, entry)
            heading_run = heading.add_run(heading_text)
            apply_word_font(heading_run, size=10, bold=True)

            in_ai_response = False
            for line in self._body_lines(entry_type, entry):
                if line == "ChatGPT回应：":
                    in_ai_response = True
                self._add_body_paragraph(doc, line, is_ai_response=in_ai_response)

        doc.save(self.path)

    def _heading_text(self, entry_type, entry):
        if entry_type == "task":
            return f"{entry['end_time']}  完成事项：{entry['title']}"
        if entry["title"]:
            return f"{entry['start_time']}  随笔：{entry['title']}"
        return f"{entry['start_time']}  随笔"

    def _body_lines(self, entry_type, entry):
        if entry_type == "task":
            lines = [
                f"计划日期：{entry['due_date']}",
                f"开始时间：{entry['start_time'] or ''}",
                f"结束时间：{entry['end_time'] or ''}",
                f"用时：{duration_text(entry['start_time'], entry['end_time'])}",
            ]
            if entry["notes"]:
                lines.append(f"备注：{entry['notes']}")
            return lines
        lines = entry["content"].splitlines() or [""]
        if entry["ai_response"]:
            lines.append("ChatGPT回应：")
            lines.extend(entry["ai_response"].splitlines())
        return lines

    def _add_body_paragraph(self, doc, text, is_ai_response=False):
        style = None
        bold = False
        left_indent = Inches(0)
        first_line_indent = Inches(0.25)
        color = None
        if text.startswith("- "):
            text = text[2:].strip()
            style = "List Bullet"
            first_line_indent = None
        elif text.startswith("> "):
            text = text[2:].strip()
            bold = True
            left_indent = Inches(0.22)
            first_line_indent = Inches(0)
            color = RGBColor(40, 40, 40)
        elif text == "ChatGPT回应：":
            bold = True
            first_line_indent = Inches(0)
        elif is_ai_response:
            first_line_indent = Inches(0)

        paragraph = doc.add_paragraph()
        if style:
            paragraph.style = style
        paragraph.paragraph_format.left_indent = left_indent
        paragraph.paragraph_format.first_line_indent = first_line_indent
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(text)
        apply_word_font(run, size=8.5, bold=bold)
        if color:
            run.font.color.rgb = color


class TaskDialog(QDialog):
    def __init__(self, parent=None, task=None):
        super().__init__(parent)
        self.setWindowTitle("编辑任务" if task else "新增任务")
        self.title_edit = QLineEdit()
        self.date_edit = QDateEdit()
        self.date_edit.setCalendarPopup(True)
        self.date_edit.setDate(QDate.currentDate())
        self.notes_edit = QPlainTextEdit()
        self.notes_edit.setPlaceholderText("备注，可留空")
        self.notes_edit.setFixedHeight(90)

        if task:
            self.title_edit.setText(task["title"])
            self.notes_edit.setPlainText(task["notes"] or "")
            self.date_edit.setDate(QDate.fromString(task["due_date"], "yyyy-MM-dd"))

        form = QFormLayout()
        form.addRow("事项", self.title_edit)
        form.addRow("计划日期", self.date_edit)
        form.addRow("备注", self.notes_edit)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout = QVBoxLayout(self)
        layout.addLayout(form)
        layout.addWidget(buttons)
        self.resize(420, 240)

    def data(self):
        return {
            "title": self.title_edit.text().strip(),
            "notes": self.notes_edit.toPlainText().strip(),
            "due_date": self.date_edit.date().toString("yyyy-MM-dd"),
        }

    def accept(self):
        if not self.title_edit.text().strip():
            QMessageBox.warning(self, "缺少事项", "请先填写事项内容。")
            return
        super().accept()


class ChatGPTResponseDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("粘贴 ChatGPT 回应")
        self.raw_edit = QPlainTextEdit()
        self.raw_edit.setPlaceholderText("把 ChatGPT 的回答粘贴到这里，保存时会自动清理空行和多余符号。")
        self.raw_edit.setMinimumHeight(240)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout = QVBoxLayout(self)
        layout.addWidget(self.raw_edit)
        layout.addWidget(buttons)
        self.resize(560, 360)

    def cleaned_text(self):
        return clean_chatgpt_response(self.raw_edit.toPlainText())

    def accept(self):
        if not self.raw_edit.toPlainText().strip():
            QMessageBox.warning(self, "缺少内容", "请先粘贴 ChatGPT 的回答。")
            return
        if not self.cleaned_text():
            QMessageBox.warning(self, "无法清理", "没有得到可写入的有效内容。")
            return
        super().accept()


class EssayDialog(QDialog):
    def __init__(self, parent=None, essay=None):
        super().__init__(parent)
        self.start_time = essay["start_time"] if essay else now_text()
        self.setWindowTitle("编辑随笔" if essay else "新随笔")

        self.time_label = QLabel(self.start_time)
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("标题，可留空")
        self.content_edit = QPlainTextEdit()
        self.content_edit.setPlaceholderText("写下今天的经历、心情、想法或突然冒出来的感悟。")
        self.content_edit.setMinimumHeight(210)
        self.ai_response_edit = QPlainTextEdit()
        self.ai_response_edit.setPlaceholderText("这里会保存清理后的 ChatGPT 回应，可手动调整。")
        self.ai_response_edit.setMinimumHeight(120)

        if essay:
            self.title_edit.setText(essay["title"] or "")
            self.content_edit.setPlainText(essay["content"] or "")
            self.ai_response_edit.setPlainText(essay["ai_response"] or "")

        form = QFormLayout()
        form.addRow("开始时间", self.time_label)
        form.addRow("标题", self.title_edit)
        form.addRow("内容", self.content_edit)
        form.addRow("ChatGPT回应", self.ai_response_edit)

        chatgpt_button = QPushButton("复制并打开 ChatGPT")
        chatgpt_button.clicked.connect(self.copy_and_open_chatgpt)
        paste_response_button = QPushButton("粘贴 ChatGPT 回应")
        paste_response_button.clicked.connect(self.paste_chatgpt_response)

        chatgpt_layout = QHBoxLayout()
        chatgpt_layout.addWidget(chatgpt_button)
        chatgpt_layout.addWidget(paste_response_button)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout = QVBoxLayout(self)
        layout.addLayout(form)
        layout.addLayout(chatgpt_layout)
        layout.addWidget(buttons)
        self.resize(620, 560)

    def data(self):
        return {
            "title": self.title_edit.text().strip(),
            "content": self.content_edit.toPlainText().strip(),
            "ai_response": clean_chatgpt_response(self.ai_response_edit.toPlainText()),
            "start_time": self.start_time,
        }

    def accept(self):
        if not self.content_edit.toPlainText().strip():
            QMessageBox.warning(self, "缺少内容", "请先写一点内容。")
            return
        super().accept()

    def copy_and_open_chatgpt(self):
        content = self.content_edit.toPlainText().strip()
        if not content:
            QMessageBox.warning(self, "缺少内容", "请先写一点内容，再打开 ChatGPT。")
            return

        title = self.title_edit.text().strip() or "无标题"
        prompt = (
            "我想让你基于你对我的了解，回应下面这篇随笔。请不要急着下结论，"
            "先理解我的感受，再帮我看见可能更平衡的角度。如果我的想法有些偏激，"
            "请温和地指出；如果只是普通记录，也可以自然回应。\n\n"
            f"开始写作时间：{self.start_time}\n"
            f"标题：{title}\n\n"
            f"随笔内容：\n{content}"
        )
        QApplication.clipboard().setText(prompt)
        webbrowser.open(CHATGPT_URL)
        QMessageBox.information(
            self,
            "已复制",
            "已经把这篇随笔整理成提示词并复制到剪贴板，同时打开了 ChatGPT。进入网页后直接粘贴发送即可。",
        )

    def paste_chatgpt_response(self):
        dialog = ChatGPTResponseDialog(self)
        clipboard_text = QApplication.clipboard().text().strip()
        if clipboard_text:
            dialog.raw_edit.setPlainText(clipboard_text)
        if dialog.exec() == QDialog.Accepted:
            self.ai_response_edit.setPlainText(dialog.cleaned_text())


class SimpleTable(QWidget):
    def __init__(self, headers, stretches, parent=None):
        super().__init__(parent)
        self.headers = headers
        self.stretches = stretches
        self.selected_id = None
        self.row_frames = []
        self.text_color = QColor("#0a0e18")

        self.setAttribute(Qt.WA_TranslucentBackground, True)
        self.setAutoFillBackground(False)

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self.header = QWidget()
        self.header.setAttribute(Qt.WA_TranslucentBackground, True)
        self.header.setAutoFillBackground(False)
        header_layout = QGridLayout(self.header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(0)
        for col, header in enumerate(headers):
            label = QLabel(header)
            label.setAlignment(Qt.AlignCenter)
            label.setMinimumHeight(30)
            label.setStyleSheet(
                "background: rgba(246, 248, 252, 175);"
                "border-right: 1px solid rgba(20, 24, 34, 42);"
                "border-bottom: 1px solid rgba(20, 24, 34, 52);"
                "padding: 4px;"
            )
            header_layout.addWidget(label, 0, col)
            header_layout.setColumnStretch(col, stretches[col])
        root.addWidget(self.header)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setFrameShape(QFrame.NoFrame)
        self.scroll.setAttribute(Qt.WA_TranslucentBackground, True)
        self.scroll.setAutoFillBackground(False)
        self.scroll.viewport().setAttribute(Qt.WA_TranslucentBackground, True)
        self.scroll.viewport().setAutoFillBackground(False)

        self.body = QWidget()
        self.body.setAttribute(Qt.WA_TranslucentBackground, True)
        self.body.setAutoFillBackground(False)
        self.body_layout = QVBoxLayout(self.body)
        self.body_layout.setContentsMargins(0, 0, 0, 0)
        self.body_layout.setSpacing(0)
        self.body_layout.addStretch(1)
        self.scroll.setWidget(self.body)
        root.addWidget(self.scroll, 1)

    def show_cell_detail(self, title, text):
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.resize(520, 420)
        layout = QVBoxLayout(dialog)
        text_edit = QPlainTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setPlainText(text)
        layout.addWidget(text_edit)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok)
        buttons.accepted.connect(dialog.accept)
        layout.addWidget(buttons)
        dialog.exec()

    def set_text_color(self, color):
        self.text_color = QColor(color)
        self.refresh_row_styles()

    def set_rows(self, rows):
        while self.body_layout.count() > 1:
            item = self.body_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()
        self.row_frames = []
        valid_ids = set()
        for row in rows:
            valid_ids.add(row["id"])
            frame = QFrame()
            frame.setProperty("row_id", row["id"])
            frame.setProperty("status", row.get("status", ""))
            frame.setCursor(Qt.PointingHandCursor)
            frame.mousePressEvent = lambda event, f=frame: self.select_row(f)
            frame.setMinimumHeight(58)
            layout = QGridLayout(frame)
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(0)
            for col, value in enumerate(row["values"]):
                label = ElidedLabel(str(value), self.headers[col], self, frame, lines=2, compact=len(str(value)) <= 24)
                label.setAlignment(Qt.AlignVCenter | (Qt.AlignCenter if col and len(str(value)) <= 22 else Qt.AlignLeft))
                label.setStyleSheet(
                    "border-right: 1px solid rgba(20, 24, 34, 38);"
                    "border-bottom: 1px solid rgba(20, 24, 34, 34);"
                    "padding: 4px;"
                )
                layout.addWidget(label, 0, col)
                layout.setColumnStretch(col, self.stretches[col])
            self.body_layout.insertWidget(self.body_layout.count() - 1, frame)
            self.row_frames.append(frame)
        if self.selected_id not in valid_ids:
            self.selected_id = None
        self.refresh_row_styles()

    def select_row(self, frame):
        self.selected_id = frame.property("row_id")
        self.refresh_row_styles()

    def refresh_row_styles(self):
        text = f"rgba({self.text_color.red()}, {self.text_color.green()}, {self.text_color.blue()}, {self.text_color.alpha()})"
        for frame in self.row_frames:
            status = frame.property("status")
            if frame.property("row_id") == self.selected_id:
                background = "rgba(47, 93, 158, 150)"
                color = "white"
            elif status == "running":
                background = "rgba(245, 158, 11, 82)"
                color = text
            elif status == "done":
                background = "rgba(34, 197, 94, 70)"
                color = text
            else:
                background = "transparent"
                color = text
            frame.setStyleSheet(f"QFrame {{ background: {background}; }} QLabel {{ color: {color}; }}")


class ElidedLabel(QLabel):
    def __init__(self, full_text, detail_title, table, row_frame, lines=2, compact=False):
        super().__init__()
        self.full_text = full_text or ""
        self.detail_title = detail_title
        self.table = table
        self.row_frame = row_frame
        self.lines = max(1, lines)
        self.compact = compact
        self.is_elided = False
        self.setMinimumHeight(28 * self.lines)
        self.setWordWrap(True)
        self.setTextInteractionFlags(Qt.NoTextInteraction)
        self.refresh_text()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.refresh_text()

    def refresh_text(self):
        display_text = self.multiline_elided_text()
        self.is_elided = display_text != self.full_text
        self.setText(display_text)
        self.setToolTip("点击查看完整内容" if self.is_elided else "")

    def multiline_elided_text(self):
        text = " ".join(self.full_text.split())
        if not text:
            return ""
        if self.compact:
            return text
        available_width = max(24, self.width() - 10)
        font_metrics = self.fontMetrics()
        words = list(text)
        lines = []
        current = ""
        index = 0
        while index < len(words) and len(lines) < self.lines:
            candidate = current + words[index]
            if font_metrics.horizontalAdvance(candidate) <= available_width or not current:
                current = candidate
                index += 1
            else:
                lines.append(current)
                current = ""
        if current and len(lines) < self.lines:
            lines.append(current)
        if index < len(words):
            if not lines:
                lines = [""]
            last = lines[-1]
            lines[-1] = font_metrics.elidedText(last + "".join(words[index:]), Qt.ElideRight, available_width)
        return "\n".join(lines)

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.table.select_row(self.row_frame)
            if self.is_elided and self.full_text.strip():
                self.table.show_cell_detail(self.detail_title, self.full_text)
                event.accept()
                return
        super().mousePressEvent(event)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.store = TaskStore(DB_PATH)
        self.journal_logger = JournalWordLogger(JOURNAL_DOCX_PATH)
        self.settings = QSettings("Codexworks", APP_NAME)
        self.tables = {}
        self.background_opacity = self.settings.value("background_opacity", 35, int)
        self.text_color = QColor(self.settings.value("text_color", "#0a0e18", str))
        self.text_opacity = self.settings.value("text_opacity", 100, int)
        self.drag_position = None
        self.resize_border = 8

        self.setWindowFlags(Qt.Window | Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground, True)
        self.setAutoFillBackground(False)
        self.setWindowOpacity(1.0)
        self.setWindowTitle(APP_DISPLAY_NAME)
        if ICON_PATH.exists():
            self.setWindowIcon(QIcon(str(ICON_PATH)))
        self.setMinimumSize(390, 560)
        self.resize(self.settings.value("width", 440, int), self.settings.value("height", 720, int))
        self.apply_visual_settings()

        self.root_panel = QWidget()
        self.root_panel.setObjectName("RootPanel")
        self.root_panel.setAttribute(Qt.WA_TranslucentBackground, True)
        self.root_panel.setAutoFillBackground(False)
        self.size_grip = QSizeGrip(self.root_panel)
        root_layout = QVBoxLayout(self.root_panel)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)
        self.setCentralWidget(self.root_panel)

        root_layout.addWidget(self.create_title_bar())

        self.tabs = QTabWidget()
        for key, title in [
            ("today", "今天"),
            ("soon", "三天提醒"),
            ("future", "未来"),
            ("done", "已完成"),
        ]:
            self.tables[key] = self.create_table()
            self.tabs.addTab(self.tables[key], title)
        self.essay_table = self.create_essay_table()
        self.tabs.addTab(self.essay_table, "随笔")
        self.configure_transparent_widget_tree(self.tabs)

        root_layout.addWidget(self.create_toolbar())
        root_layout.addWidget(self.tabs, 1)
        self.apply_always_on_top(self.settings.value("always_on_top", False, bool))
        self.refresh()

    def create_title_bar(self):
        bar = QWidget()
        bar.setObjectName("TitleBar")
        bar.setFixedHeight(34)
        layout = QHBoxLayout(bar)
        layout.setContentsMargins(10, 4, 8, 4)
        layout.setSpacing(6)

        title = QLabel(APP_DISPLAY_NAME)
        if ICON_PATH.exists():
            icon_label = QLabel()
            icon_label.setPixmap(QIcon(str(ICON_PATH)).pixmap(16, 16))
            layout.addWidget(icon_label)
        layout.addWidget(title)
        layout.addStretch(1)

        settings_button = QPushButton("设置")
        settings_button.setObjectName("TitleButton")
        settings_button.setMenu(self.create_settings_menu())
        layout.addWidget(settings_button)

        minimize_button = QPushButton("—")
        minimize_button.setObjectName("TitleButton")
        minimize_button.clicked.connect(self.showMinimized)
        layout.addWidget(minimize_button)

        close_button = QPushButton("×")
        close_button.setObjectName("TitleButton")
        close_button.clicked.connect(self.close)
        layout.addWidget(close_button)

        bar.mousePressEvent = self.title_mouse_press_event
        bar.mouseMoveEvent = self.title_mouse_move_event
        bar.mouseReleaseEvent = self.title_mouse_release_event
        return bar

    def create_settings_menu(self):
        settings_menu = QMenu("设置", self)

        opacity_action = QAction("背景透明度", self)
        opacity_action.triggered.connect(self.choose_background_opacity)
        settings_menu.addAction(opacity_action)

        text_color_action = QAction("字体颜色", self)
        text_color_action.triggered.connect(self.choose_text_color)
        settings_menu.addAction(text_color_action)

        text_opacity_action = QAction("字体透明度", self)
        text_opacity_action.triggered.connect(self.choose_text_opacity)
        settings_menu.addAction(text_opacity_action)

        window_size_action = QAction("窗口大小", self)
        window_size_action.triggered.connect(self.choose_window_size)
        settings_menu.addAction(window_size_action)

        settings_menu.addSeparator()

        reset_visual_action = QAction("恢复默认外观", self)
        reset_visual_action.triggered.connect(self.reset_visual_settings)
        settings_menu.addAction(reset_visual_action)

        return settings_menu

    def title_mouse_press_event(self, event):
        if event.button() == Qt.LeftButton:
            self.drag_position = event.globalPosition().toPoint() - self.frameGeometry().topLeft()
            event.accept()

    def title_mouse_move_event(self, event):
        if self.drag_position is not None and event.buttons() & Qt.LeftButton:
            self.move(event.globalPosition().toPoint() - self.drag_position)
            event.accept()

    def title_mouse_release_event(self, event):
        self.drag_position = None
        event.accept()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if hasattr(self, "size_grip"):
            size = self.size_grip.sizeHint()
            self.size_grip.move(self.root_panel.width() - size.width() - 4, self.root_panel.height() - size.height() - 4)

    def nativeEvent(self, event_type, message):
        if sys.platform != "win32":
            return super().nativeEvent(event_type, message)
        try:
            msg = ctypes.wintypes.MSG.from_address(int(message))
        except Exception:
            return super().nativeEvent(event_type, message)

        wm_nchittest = 0x0084
        if msg.message != wm_nchittest:
            return super().nativeEvent(event_type, message)

        x = ctypes.c_short(msg.lParam & 0xFFFF).value
        y = ctypes.c_short((msg.lParam >> 16) & 0xFFFF).value
        rect = self.frameGeometry()
        left = x - rect.left() <= self.resize_border
        right = rect.right() - x <= self.resize_border
        top = y - rect.top() <= self.resize_border
        bottom = rect.bottom() - y <= self.resize_border

        htleft, htright, httop, htbottom = 10, 11, 12, 15
        httopleft, httopright, htbottomleft, htbottomright = 13, 14, 16, 17
        if top and left:
            return True, httopleft
        if top and right:
            return True, httopright
        if bottom and left:
            return True, htbottomleft
        if bottom and right:
            return True, htbottomright
        if left:
            return True, htleft
        if right:
            return True, htright
        if top:
            return True, httop
        if bottom:
            return True, htbottom
        return super().nativeEvent(event_type, message)

    def create_toolbar(self):
        toolbar = QToolBar()
        toolbar.setMovable(False)

        add_action = QAction("新增", self)
        add_action.triggered.connect(self.add_task)
        toolbar.addAction(add_action)

        edit_action = QAction("编辑", self)
        edit_action.triggered.connect(self.edit_task)
        toolbar.addAction(edit_action)

        delete_action = QAction("删除", self)
        delete_action.triggered.connect(self.delete_task)
        toolbar.addAction(delete_action)

        start_action = QAction("开始", self)
        start_action.triggered.connect(self.start_task)
        toolbar.addAction(start_action)

        finish_action = QAction("结束", self)
        finish_action.triggered.connect(self.finish_task)
        toolbar.addAction(finish_action)

        toolbar.addSeparator()

        essay_action = QAction("新随笔", self)
        essay_action.triggered.connect(self.add_essay)
        toolbar.addAction(essay_action)

        edit_essay_action = QAction("编辑随笔", self)
        edit_essay_action.triggered.connect(self.edit_essay)
        toolbar.addAction(edit_essay_action)

        delete_essay_action = QAction("删除随笔", self)
        delete_essay_action.triggered.connect(self.delete_essay)
        toolbar.addAction(delete_essay_action)

        toolbar.addSeparator()

        open_doc_action = QAction("打开生活记录", self)
        open_doc_action.triggered.connect(self.open_journal_log)
        toolbar.addAction(open_doc_action)

        options = QWidget()
        option_layout = QHBoxLayout(options)
        option_layout.setContentsMargins(8, 0, 0, 0)

        self.top_check = QCheckBox("置顶")
        self.top_check.setChecked(self.settings.value("always_on_top", False, bool))
        self.top_check.toggled.connect(self.apply_always_on_top)
        option_layout.addWidget(self.top_check)

        self.startup_check = QCheckBox("开机自启")
        self.startup_check.setChecked(is_startup_enabled())
        self.startup_check.toggled.connect(self.toggle_startup)
        option_layout.addWidget(self.startup_check)

        toolbar.addWidget(options)
        return toolbar

    def apply_visual_settings(self):
        self.background_opacity = max(8, min(90, int(self.background_opacity)))
        self.text_opacity = max(20, min(100, int(self.text_opacity)))
        if not self.text_color.isValid():
            self.text_color = QColor("#0a0e18")
        self.setWindowOpacity(1.0)
        self.setStyleSheet(build_transparent_style(self.text_color, self.text_opacity, self.background_opacity))
        self.settings.setValue("background_opacity", self.background_opacity)
        self.settings.setValue("text_color", self.text_color.name())
        self.settings.setValue("text_opacity", self.text_opacity)
        self.settings.remove("background_color")
        if hasattr(self, "root_panel"):
            self.configure_transparent_widget_tree(self.root_panel)
        if hasattr(self, "tabs"):
            self.configure_transparent_widget_tree(self.tabs)
        for table in list(self.tables.values()) + ([self.essay_table] if hasattr(self, "essay_table") else []):
            if isinstance(table, SimpleTable):
                table.set_text_color(self.current_text_qcolor())
                self.configure_transparent_widget_tree(table)
            else:
                self.configure_transparent_table(table)

    def choose_background_opacity(self):
        value, ok = QInputDialog.getInt(
            self,
            "透明度",
            "背景不透明度（8=很透明，90=更实）：",
            self.background_opacity,
            8,
            90,
            1,
        )
        if ok:
            self.background_opacity = value
            self.apply_visual_settings()
            self.refresh()

    def choose_text_color(self):
        color = QColorDialog.getColor(self.text_color, self, "选择字体颜色")
        if color.isValid():
            self.text_color = color
            self.apply_visual_settings()
            self.refresh()

    def choose_text_opacity(self):
        value, ok = QInputDialog.getInt(
            self,
            "字体透明度",
            "字体不透明度（20=较淡，100=实心）：",
            self.text_opacity,
            20,
            100,
            1,
        )
        if ok:
            self.text_opacity = value
            self.apply_visual_settings()
            self.refresh()

    def choose_window_size(self):
        width, ok = QInputDialog.getInt(
            self,
            "窗口大小",
            "窗口宽度：",
            self.width(),
            self.minimumWidth(),
            2400,
            10,
        )
        if not ok:
            return
        height, ok = QInputDialog.getInt(
            self,
            "窗口大小",
            "窗口高度：",
            self.height(),
            self.minimumHeight(),
            1800,
            10,
        )
        if not ok:
            return
        self.resize(width, height)
        self.settings.setValue("width", width)
        self.settings.setValue("height", height)

    def reset_visual_settings(self):
        self.background_opacity = 35
        self.text_color = QColor("#0a0e18")
        self.text_opacity = 100
        self.apply_visual_settings()
        self.refresh()

    def create_table(self):
        return SimpleTable(["事项", "计划日期", "状态", "开始", "结束", "备注"], [5, 1, 1, 1, 1, 5])

    def create_essay_table(self):
        return SimpleTable(["开始时间", "标题", "内容"], [2, 2, 6])
    def configure_transparent_table(self, table):
        self.configure_transparent_widget_tree(table)
        table.setFrameShape(QFrame.NoFrame)
        table.setAutoFillBackground(False)
        table.setAlternatingRowColors(False)
        table.setAttribute(Qt.WA_TranslucentBackground, True)
        table.setAttribute(Qt.WA_NoSystemBackground, True)
        table.setAttribute(Qt.WA_OpaquePaintEvent, False)
        table.viewport().setAutoFillBackground(False)
        table.viewport().setAttribute(Qt.WA_TranslucentBackground, True)
        table.viewport().setAttribute(Qt.WA_NoSystemBackground, True)
        table.viewport().setAttribute(Qt.WA_OpaquePaintEvent, False)
        palette = table.palette()
        transparent = QColor(0, 0, 0, 0)
        palette.setColor(QPalette.Base, transparent)
        palette.setColor(QPalette.Window, transparent)
        palette.setColor(QPalette.Text, self.current_text_qcolor())
        palette.setBrush(QPalette.Base, QBrush(Qt.NoBrush))
        palette.setBrush(QPalette.Window, QBrush(Qt.NoBrush))
        table.setPalette(palette)
        table.viewport().setPalette(palette)

    def current_text_qcolor(self):
        color = QColor(self.text_color)
        color.setAlpha(max(30, min(255, int(255 * self.text_opacity / 100))))
        return color

    def configure_transparent_widget_tree(self, widget):
        transparent = QColor(0, 0, 0, 0)
        stack = [widget]
        while stack:
            current = stack.pop()
            current.setAutoFillBackground(False)
            current.setAttribute(Qt.WA_TranslucentBackground, True)
            current.setAttribute(Qt.WA_NoSystemBackground, True)
            current.setAttribute(Qt.WA_OpaquePaintEvent, False)
            palette = current.palette()
            palette.setColor(QPalette.Window, transparent)
            palette.setColor(QPalette.Base, transparent)
            palette.setBrush(QPalette.Window, QBrush(Qt.NoBrush))
            palette.setBrush(QPalette.Base, QBrush(Qt.NoBrush))
            current.setPalette(palette)
            stack.extend(current.findChildren(QWidget, options=Qt.FindDirectChildrenOnly))

    def selected_task_id(self):
        table = self.tabs.currentWidget()
        if table not in self.tables.values():
            QMessageBox.information(self, "当前不是任务页", "请先切换到任务列表再操作任务。")
            return None
        if table.selected_id is None:
            QMessageBox.information(self, "未选择任务", "请先选择一条任务。")
            return None
        return table.selected_id

    def selected_essay_id(self):
        if self.tabs.currentWidget() is not self.essay_table:
            self.tabs.setCurrentWidget(self.essay_table)
        if self.essay_table.selected_id is None:
            QMessageBox.information(self, "未选择随笔", "请先选择一篇随笔。")
            return None
        return self.essay_table.selected_id
    def add_task(self):
        dialog = TaskDialog(self)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.data()
            self.store.add_task(data["title"], data["notes"], data["due_date"])
            self.refresh()

    def edit_task(self):
        task_id = self.selected_task_id()
        if task_id is None:
            return
        task = self.store.get_task(task_id)
        dialog = TaskDialog(self, task)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.data()
            self.store.update_task(task_id, data["title"], data["notes"], data["due_date"])
            self.refresh()

    def delete_task(self):
        task_id = self.selected_task_id()
        if task_id is None:
            return
        if QMessageBox.question(self, "确认删除", "确定删除这条任务吗？") == QMessageBox.Yes:
            self.store.delete_task(task_id)
            self.refresh()

    def start_task(self):
        task_id = self.selected_task_id()
        if task_id is None:
            return
        self.store.start_task(task_id)
        self.refresh()

    def finish_task(self):
        task_id = self.selected_task_id()
        if task_id is None:
            return
        task = self.store.complete_task(task_id)
        if task and not task["logged_at"]:
            self.store.mark_logged(task_id)
        self.sync_journal_log()
        self.refresh()
        QMessageBox.information(self, "已完成", f"任务已完成，并写入生活记录：\n{JOURNAL_DOCX_PATH}")

    def add_essay(self):
        dialog = EssayDialog(self)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.data()
            self.store.add_essay(data["title"], data["content"], data["ai_response"], data["start_time"])
            self.sync_journal_log()
            self.refresh()

    def edit_essay(self):
        essay_id = self.selected_essay_id()
        if essay_id is None:
            return
        essay = self.store.get_essay(essay_id)
        dialog = EssayDialog(self, essay)
        if dialog.exec() == QDialog.Accepted:
            data = dialog.data()
            self.store.update_essay(essay_id, data["title"], data["content"], data["ai_response"])
            self.sync_journal_log()
            self.refresh()

    def delete_essay(self):
        essay_id = self.selected_essay_id()
        if essay_id is None:
            return
        if QMessageBox.question(self, "确认删除", "确定删除这篇随笔吗？") == QMessageBox.Yes:
            self.store.delete_essay(essay_id)
            self.sync_journal_log()
            self.refresh()

    def open_journal_log(self):
        if not JOURNAL_DOCX_PATH.exists():
            self.sync_journal_log()
        if not JOURNAL_DOCX_PATH.exists():
            QMessageBox.information(self, "暂无记录", "还没有完成事项或保存随笔。")
            return
        os.startfile(JOURNAL_DOCX_PATH)

    def sync_journal_log(self):
        tasks = self.store.list_tasks()
        essays = self.store.list_essays()
        has_completed_tasks = any(task["status"] == "done" and task["end_time"] for task in tasks)
        if has_completed_tasks or essays:
            self.journal_logger.rewrite(tasks, essays)
        elif JOURNAL_DOCX_PATH.exists():
            JOURNAL_DOCX_PATH.unlink()

    def apply_always_on_top(self, checked):
        self.settings.setValue("always_on_top", checked)
        flags = self.windowFlags()
        if checked:
            flags |= Qt.WindowStaysOnTopHint
        else:
            flags &= ~Qt.WindowStaysOnTopHint
        self.setWindowFlags(flags)
        self.show()

    def toggle_startup(self, checked):
        try:
            set_startup_enabled(checked)
        except Exception as exc:
            self.startup_check.blockSignals(True)
            self.startup_check.setChecked(not checked)
            self.startup_check.blockSignals(False)
            QMessageBox.warning(self, "设置失败", f"无法修改开机自启动：\n{exc}")

    def refresh(self):
        buckets = {"today": [], "soon": [], "future": [], "done": []}
        today = date.today()
        soon_end = today + timedelta(days=3)
        for task in self.store.list_tasks():
            if task["status"] == "done":
                buckets["done"].append(task)
                continue
            due = parse_date(task["due_date"])
            if due == today:
                buckets["today"].append(task)
            elif today < due <= soon_end:
                buckets["soon"].append(task)
            else:
                buckets["future"].append(task)

        for key, rows in buckets.items():
            self.populate_table(self.tables[key], rows)
            index = list(self.tables).index(key)
            self.tabs.setTabText(index, f"{self.tabs.tabText(index).split(' ')[0]} ({len(rows)})")
        essays = self.store.list_essays()
        self.populate_essay_table(essays)
        self.tabs.setTabText(self.tabs.indexOf(self.essay_table), f"随笔 ({len(essays)})")

    def populate_table(self, table, rows):
        table.set_text_color(self.current_text_qcolor())
        table.set_rows([
            {
                "id": task["id"],
                "status": task["status"],
                "values": [
                    task["title"],
                    task["due_date"],
                    status_label(task["status"]),
                    task["start_time"] or "",
                    task["end_time"] or "",
                    task["notes"] or "",
                ],
            }
            for task in rows
        ])

    def populate_essay_table(self, rows):
        self.essay_table.set_text_color(self.current_text_qcolor())
        self.essay_table.set_rows([
            {
                "id": essay["id"],
                "status": "",
                "values": [
                    essay["start_time"],
                    essay["title"] or "无标题",
                    " ".join((essay["content"] or "").split()),
                ],
            }
            for essay in rows
        ])
    def closeEvent(self, event):
        self.settings.setValue("width", self.width())
        self.settings.setValue("height", self.height())
        self.store.close()
        super().closeEvent(event)


def status_label(status):
    return {
        "pending": "未开始",
        "running": "进行中",
        "done": "已完成",
    }.get(status, status)


def startup_command():
    script = PROJECT_DIR / "app.py"
    pythonw = Path(sys.executable).with_name("pythonw.exe")
    exe = pythonw if pythonw.exists() else Path(sys.executable)
    return f'"{exe}" "{script}"'


def is_startup_enabled():
    if sys.platform != "win32":
        return False
    try:
        import winreg

        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, RUN_KEY, 0, winreg.KEY_READ) as key:
            value, _ = winreg.QueryValueEx(key, APP_NAME)
        return value == startup_command()
    except FileNotFoundError:
        return False
    except OSError:
        return False


def set_startup_enabled(enabled):
    if sys.platform != "win32":
        raise RuntimeError("开机自启动目前只支持 Windows。")
    import winreg

    with winreg.OpenKey(winreg.HKEY_CURRENT_USER, RUN_KEY, 0, winreg.KEY_SET_VALUE) as key:
        if enabled:
            winreg.SetValueEx(key, APP_NAME, 0, winreg.REG_SZ, startup_command())
        else:
            try:
                winreg.DeleteValue(key, APP_NAME)
            except FileNotFoundError:
                pass


def main():
    set_windows_app_id()
    app = QApplication(sys.argv)
    app.setApplicationName(APP_DISPLAY_NAME)
    app.setWindowIcon(QIcon(str(ICON_PATH)))
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


def set_windows_app_id():
    if sys.platform != "win32":
        return
    try:
        import ctypes

        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("Codexworks.LinsNotebook")
    except Exception:
        pass


if __name__ == "__main__":
    main()
